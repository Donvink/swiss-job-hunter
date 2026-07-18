"""
Text extraction for uploaded resume files (PDF/DOCX/TXT).
"""
from __future__ import annotations

import io

from db.models import FileFormat


def extract_text(file_bytes: bytes, file_format: FileFormat) -> str:
    """Extract plain text from an uploaded resume file."""
    if file_format == FileFormat.PDF:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if file_format == FileFormat.DOCX:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    return file_bytes.decode("utf-8", errors="replace")
